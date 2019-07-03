from docx import *
import win32com.client
import os
import sys
import docx
import getopt
from docx.shared import Inches
from docxtpl import DocxTemplate
import jinja2
#from docx import Document

word_docuemnt_filename = ''
jira_issue = ''

def main(argv):

    try:
        opts, args = getopt.getopt(argv, "i:f:f:", ["issue=", "filename="])
        print('ARGS:')

        for opt, arg in opts:
            print(opt, arg)

    except getopt.GetoptError:
        print('ERROR: Word.py -f <word_docuemnt_filename> -i <jira_issue>')
        sys.exit(2)


    print('PARSING ARGS:')

    for opt, arg in opts:
        if opt == '-h':
            print('Word.py -f <word_docuemnt_filename> -i <jira_issue>')
            sys.exit()
        else:
            print(opt, arg)
            if opt in ("-i", "--issue"):
                jira_issue = arg
            if opt in ("-f", "--filename"):
                word_docuemnt_filename = arg

    if not(os.path.isfile(word_docuemnt_filename)):
        new_document(word_docuemnt_filename)

    insert_Issue(jira_issue, word_docuemnt_filename)

def insert_Issue (issue, jira_body, jira_created, jira_status, word_docuemnt_filename):
    #jira_issue, jira_body, issue.fields.created, issue.fields.status.name, word_docuemnt_filenam
    print("Openning Document...")
    document = docx.Document(word_docuemnt_filename)

    print("Adding issue: " + issue)
    document.add_heading(issue, level=1)
    document.add_paragraph(jira_body, style='Intense Quote')
    document.add_paragraph('Created: ' + jira_created, style='Intense Quote')
    document.add_paragraph('Status: ' + jira_status, style='Intense Quote')
    #document.add_heading('Document Title: ' + issue, 0)

    print('Saving...')
    document.save(word_docuemnt_filename)

    print('Replacing...')
    replace(word_docuemnt_filename)

def new_document (new_filename):
    print('CREATING NEW DOCUMENT:')
    document = docx.Document('test.docx')

    document.add_heading('REQUESRMENTS:', 0)

    print('Saving as: ' + new_filename)
    document.save(new_filename)
    print('CREATING NEW DOCUMENT: DONE!')

def updateTOC (word_docuemnt_filename):
    print('Updating TOC...')
    dir_path = os.path.dirname(os.path.realpath(__file__))
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(dir_path + "\\" +  word_docuemnt_filename)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()

def multiply_by(value, by):
   return value * by

def replace (word_docuemnt_filename):
    doc = DocxTemplate(word_docuemnt_filename)
    context = {'price_dollars': 5.00}
    jinja_env = jinja2.Environment()
    jinja_env.filters['multiply_by'] = multiply_by
    doc.render(context, jinja_env)
    doc.save(word_docuemnt_filename)

if __name__ == "__main__":
   main(sys.argv[1:])